"""Generate the FLUX.2-klein PTL/OpenVINO BKM as a Word (.docx) document.

The environment header, results table and per-case details are parsed from the
REAL benchmark output (``benchmark_outputs_i2i/benchmark_results.csv``) so the
document reflects measured data rather than hard-coded numbers. The narrative
(enablement steps, methodology, troubleshooting) is BKM prose.

Usage:
    python make_bkm_docx.py \
        --results benchmark_outputs_i2i/benchmark_results.csv \
        --images-dir benchmark_outputs_i2i/images \
        --output FLUX2-klein_PTL_OpenVINO_BKM.docx
"""

from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# --------------------------------------------------------------------------- #
# Parse the real benchmark CSV
# --------------------------------------------------------------------------- #
def parse_results_csv(path: Path):
    """Return (env dict, table headers, table rows, notes list) from the CSV."""
    env: dict = {}
    headers: list = []
    rows: list = []
    notes: list = []
    section = None

    with path.open(encoding="utf-8") as fh:
        for record in csv.reader(fh):
            if not record:
                continue
            first = record[0]
            if first == "# Environment":
                section = "env"
                continue
            if first == "# Notes":
                section = "notes"
                continue
            if first == "Resolution":  # results table header row
                headers = record
                section = "table"
                continue
            if section == "env" and first.startswith("# "):
                env[first[2:]] = record[1] if len(record) > 1 else ""
            elif section == "table":
                rows.append(record)
            elif section == "notes" and first.startswith("# "):
                notes.append(first[2:])
    return env, headers, rows, notes


def build_summary(notes: list):
    """Derive an executive summary (tier, resolution, IR read, e2e/image, peak
    RAM) from the per-case notes so no numbers are hard-coded."""
    pattern = re.compile(
        r"(\w+) tier @ (.+?):.*?pipeline IR read = ([0-9.]+s).*?"
        r"end-to-end inference = ([0-9.]+s) per image; Peak RAM ([0-9.]+G)"
    )
    summary = []
    for note in notes:
        m = pattern.search(note)
        if m:
            tier, resolution, ir_read, e2e, peak = m.groups()
            summary.append([tier.upper(), resolution, ir_read, e2e, peak])
    return summary


# --------------------------------------------------------------------------- #
# docx helpers
# --------------------------------------------------------------------------- #
def shade_paragraph(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code_block(doc: Document, code: str):
    """Add a monospace, light-gray shaded code block."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(6)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    shade_paragraph(para)
    return para


def add_table(doc: Document, headers: list, rows: list, style: str = "Light Grid Accent 1"):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = style
    except Exception:
        table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
    return table


def add_caption(doc: Document, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_image_grid(doc: Document, images: list, images_dir: Path):
    """Place (filename, caption) images in a 2-column table."""
    available = [(images_dir / name, cap) for name, cap in images if (images_dir / name).exists()]
    if not available:
        doc.add_paragraph("(Generated sample images were not found at doc-build time.)")
        return
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx in range(0, len(available), 2):
        cells = table.add_row().cells
        for col, (path, caption) in enumerate(available[idx : idx + 2]):
            cell = cells[col]
            p_img = cell.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.add_run().add_picture(str(path), width=Inches(2.7))
            except Exception:
                p_img.add_run(f"[could not embed {path.name}]")
            p_cap = cell.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cap.add_run(caption)
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_io_gallery(doc: Document, pairs: list, input_dir: Path, output_dir: Path,
                   img_width=Inches(2.0)):
    """Render Input -> FP16 output -> INT4 output rows in a 3-column table."""
    table = doc.add_table(rows=1, cols=3)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(["Input photo", "FP16 output", "INT4 output"]):
        run = table.rows[0].cells[i].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
    for spec in pairs:
        cells = table.add_row().cells
        entries = [
            (input_dir / spec["input"], spec["res"]),
            (output_dir / spec["fp16"], f'{spec["res"]} - FP16'),
            (output_dir / spec["int4"], f'{spec["res"]} - INT4'),
        ]
        for col, (path, caption) in enumerate(entries):
            cell = cells[col]
            p_img = cell.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if path.exists():
                try:
                    p_img.add_run().add_picture(str(path), width=img_width)
                except Exception:
                    p_img.add_run(f"[could not embed {path.name}]")
            else:
                p_img.add_run(f"[missing {path.name}]")
            p_cap = cell.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p_cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return table


# --------------------------------------------------------------------------- #
# Build the document
# --------------------------------------------------------------------------- #
def build_document(env, headers, rows, notes, summary, images_dir: Path,
                   input_images_dir: Path, prompt: str, output: Path):
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---- Title ----------------------------------------------------------- #
    title = doc.add_heading(
        "BKM: Enabling FLUX.2-klein Image Generation on Intel Panther Lake (PTL) with OpenVINO",
        level=0,
    )
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        "Best Known Method for exporting, running and benchmarking "
        "black-forest-labs/FLUX.2-klein-4B via Optimum-Intel + OpenVINO on the Intel GPU."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ---- 1. Overview ----------------------------------------------------- #
    doc.add_heading("1. Overview & Scope", level=1)
    doc.add_paragraph(
        "This BKM describes how to deploy and benchmark the FLUX.2-klein-4B model on the Intel "
        "Panther Lake (PTL) platform using OpenVINO and Optimum-Intel."
    )

    # ---- 2. Platform & environment --------------------------------------- #
    doc.add_heading("2. Platform & Software Environment", level=1)
    doc.add_paragraph(
        "All results in this BKM were produced on the following configuration (captured "
        "automatically by the benchmark harness):"
    )
    # Timestamp is test-automation log residue with no guidance value -> omit.
    env_rows = [[k, v] for k, v in env.items() if k != "Timestamp"]
    add_table(doc, ["Item", "Value"], env_rows)

    # ---- 3. Prerequisites & install -------------------------------------- #
    doc.add_heading("3. Prerequisites & Installation", level=1)
    doc.add_paragraph(
        "Use a dedicated Python 3.10 virtual environment. Install Optimum-Intel with the OpenVINO "
        "and NNCF extras plus the diffusion/imaging dependencies:"
    )
    add_code_block(
        doc,
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate            # Windows PowerShell\n"
        "pip install \"optimum-intel[openvino,nncf]\" diffusers pillow psutil",
    )
    doc.add_paragraph("Confirm OpenVINO detects the Intel GPU (PTL iGPU):")
    add_code_block(
        doc,
        'python -c "import openvino as ov; print(ov.Core().available_devices)"\n'
        "# Expect a list that includes 'GPU' (and typically 'CPU', 'NPU').",
    )

    # ---- 4. Enable the model: export ------------------------------------- #
    doc.add_heading("4. Enable the Model — Export to OpenVINO IR", level=1)
    doc.add_paragraph(
        "FLUX.2-klein is a diffusers pipeline; each component (text encoder, transformer, VAE) is "
        "converted to OpenVINO IR at a chosen weight precision. This BKM uses two tiers:"
    )
    doc.add_paragraph("FP16 tier (Text Encoder / Transformer / VAE all FP16):")
    add_code_block(
        doc,
        "optimum-cli export openvino \\\n"
        "  --model black-forest-labs/FLUX.2-klein-4B \\\n"
        "  --weight-format fp16 \\\n"
        "  flux2-klein-fp16",
    )
    doc.add_paragraph("INT4 tier (Text Encoder + Transformer INT4, VAE kept FP16):")
    add_code_block(
        doc,
        "optimum-cli export openvino \\\n"
        "  --model black-forest-labs/FLUX.2-klein-4B \\\n"
        "  --weight-format int4 \\\n"
        "  flux2-klein-int4",
    )
    doc.add_paragraph("Tips:")
    for tip in [
        "Export once per precision tier and reuse the directory for every resolution — do not "
        "re-export on each run (resolution is chosen at inference time).",
        "Models larger than ~1 GB default to INT8 weight compression; always pass "
        "--weight-format fp16/int4 to get a well-defined precision tier.",
        "Set an HF_TOKEN environment variable before exporting for faster, rate-limit-free "
        "downloads from the Hugging Face Hub.",
    ]:
        doc.add_paragraph(tip, style="List Bullet")

    # ---- 5. Run image generation on PTL ---------------------------------- #
    doc.add_heading("5. Run Image Generation on PTL (Intel GPU)", level=1)
    doc.add_paragraph(
        "Load the exported model with device=\"GPU\" so inference runs on the PTL Intel GPU. Running "
        "on the GPU is essential — the same pipeline on CPU is dramatically slower. The helper script "
        "flux2_klein_inference.py wraps both modes."
    )
    doc.add_paragraph("Text-to-image:")
    add_code_block(
        doc,
        "python flux2_klein_inference.py text2image \\\n"
        "  --model-dir flux2-klein-fp16 \\\n"
        "  --prompt \"Astronaut on the moon\" \\\n"
        "  --height 512 --width 512 --steps 4 --device GPU",
    )
    doc.add_paragraph("Image-to-image (instruction edit, e.g. \"remove cloud from the photo\"):")
    add_code_block(
        doc,
        "python flux2_klein_inference.py image2image \\\n"
        "  --model-id flux2-klein-fp16 \\\n"
        "  --prompt \"remove cloud from the photo\" \\\n"
        "  --images images/512x512.jpg images/1024x1024.jpg \\\n"
        "  --steps 4 --device GPU",
    )
    doc.add_paragraph(
        "Programmatic equivalent (same pipeline the benchmark reuses):"
    )
    add_code_block(
        doc,
        "from optimum.intel import OVDiffusionPipeline\n"
        "pipe = OVDiffusionPipeline.from_pretrained(\"flux2-klein-fp16\", device=\"GPU\")\n"
        "image = pipe(prompt=\"Astronaut on the moon\", height=512, width=512,\n"
        "             num_inference_steps=4, output_type=\"pil\").images[0]\n"
        "image.save(\"out.png\")",
    )

    # ---- 6. Benchmark methodology ---------------------------------------- #
    doc.add_heading("6. Benchmark Methodology", level=1)
    doc.add_paragraph(
        "Two harness scripts reuse the same OpenVINO pipeline: flux2_klein_benchmark.py "
        "(text-to-image) and flux2_klein_benchmark_i2i.py (image editing). For each precision tier "
        "and resolution (512x512 and 1024x1024, 4 steps), the harness runs 3 warmup iterations "
        "(not timed) followed by 5 measured iterations and reports the average."
    )
    doc.add_paragraph("Metric definitions:")
    add_table(
        doc,
        ["Metric", "Definition"],
        [
            ["Load Time", "OpenVINO compile_model time for that component on the target device (shared IR read reported per case in notes)."],
            ["Infer Time", "Per-component forward time during measured runs only, averaged per image (Transformer covers all denoising steps)."],
            ["RAM After Load", "Cumulative process RSS captured right after that component is compiled."],
            ["Peak RAM", "Process-level peak RSS during the whole case (identical across a case's rows)."],
        ],
    )
    doc.add_paragraph("Run the benchmarks (reusing the already-exported model dirs, GPU):")
    add_code_block(
        doc,
        "python flux2_klein_benchmark_i2i.py --device GPU \\\n"
        "  --fp16-model-dir flux2-klein-fp16 --int4-model-dir flux2-klein-int4 \\\n"
        "  --resolutions 512,1024 --steps 4 --warmup 3 --runs 5 \\\n"
        "  --output-dir benchmark_outputs_i2i\n"
        "# Text-to-image variant: python flux2_klein_benchmark.py (same flags)",
    )

    # ---- 7. Results ------------------------------------------------------ #
    doc.add_heading("7. Performance Benchmark Results", level=1)
    doc.add_paragraph(
        "Workload: image editing (image-to-image), prompt \"remove cloud from the photo\", 4 steps, "
        "device=GPU. Values are measured on the platform in Section 2."
    )
    if summary:
        doc.add_heading("7.1 Executive summary (per image)", level=2)
        add_table(
            doc,
            ["Tier", "Resolution", "Pipeline load (IR)", "End-to-end / image", "Peak RAM"],
            summary,
        )
        doc.add_paragraph(
            "INT4 weight compression roughly halves end-to-end latency versus FP16 at both "
            "resolutions while cutting resident memory; the Transformer dominates inference time, "
            "whereas text encoding and VAE decoding are minor contributors.",
        )

    doc.add_heading("7.2 Per-component detail", level=2)
    add_table(doc, headers, rows)

    # ---- 8. Sample inputs & outputs -------------------------------------- #
    doc.add_heading("8. Sample Inputs & Generated Images (PTL GPU)", level=1)
    doc.add_paragraph(
        "All samples were produced by the image-editing (image-to-image) workflow on the PTL Intel "
        "GPU. Each input photo is resized to the target resolution before editing."
    )
    prompt_para = doc.add_paragraph()
    prompt_para.add_run("Editing prompt: ").bold = True
    prompt_para.add_run(f'\u201c{prompt}\u201d').italic = True
    add_io_gallery(
        doc,
        [
            {
                "res": "512x512",
                "input": "512x512.jpg",
                "fp16": "flux2_i2i_fp16_512x512_4steps.png",
                "int4": "flux2_i2i_int4_512x512_4steps.png",
            },
            {
                "res": "1024x1024",
                "input": "1024x1024.jpg",
                "fp16": "flux2_i2i_fp16_1024x1024_4steps.png",
                "int4": "flux2_i2i_int4_1024x1024_4steps.png",
            },
        ],
        input_images_dir,
        images_dir,
    )

    # ---- 9. Troubleshooting ---------------------------------------------- #
    doc.add_heading("9. Troubleshooting", level=1)
    for item in [
        "\"scaling_factor attribute is missing from the VAE ... re-export\": benign; the pipeline "
        "falls back to a default. Re-export with matching optimum-intel + diffusers versions to silence it.",
        "TracerWarnings during export are expected (torch tracing) and do not affect runtime performance.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ---- 10. Appendix ---------------------------------------------------- #
    doc.add_heading("10. Appendix — Files", level=1)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["flux2_klein_inference.py", "Run text-to-image / image-to-image generation on CPU or GPU."],
            ["flux2_klein_benchmark.py", "Text-to-image performance benchmark harness."],
            ["flux2_klein_benchmark_i2i.py", "Image-to-image performance benchmark harness."],
            ["requirements.txt", "Python dependencies (optimum-intel[openvino,nncf], diffusers, pillow, psutil, ...)."],
            ["benchmark_outputs_i2i/benchmark_results.{csv,md}", "Machine-readable + Markdown benchmark results."],
        ],
    )

    doc.save(str(output))
    print(f"Wrote {output}")


def main():
    parser = argparse.ArgumentParser(description="Build the FLUX.2-klein PTL BKM Word document.")
    parser.add_argument("--results", default="benchmark_outputs_i2i/benchmark_results.csv")
    parser.add_argument("--images-dir", default="benchmark_outputs_i2i/images")
    parser.add_argument("--input-images-dir", default="images")
    parser.add_argument("--prompt", default="remove cloud from the photo")
    parser.add_argument("--output", default="FLUX2-klein_PTL_OpenVINO_BKM.docx")
    args = parser.parse_args()

    env, headers, rows, notes = parse_results_csv(Path(args.results))
    summary = build_summary(notes)
    build_document(env, headers, rows, notes, summary, Path(args.images_dir),
                   Path(args.input_images_dir), args.prompt, Path(args.output))


if __name__ == "__main__":
    main()
